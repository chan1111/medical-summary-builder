"""ReportAgent — build the final .docx output.

- Default mode: populate the Word template using {{PLACEHOLDER}} tokens.
- Custom layout mode: build a fresh document with user-defined table columns
  (layout_instruction from IntentAgent → AnalysisAgent custom layout call).
"""

from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path

logger = logging.getLogger(__name__)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from rich.console import Console

from .base import BaseAgent
from .analysis_agent import apply_custom_layout, _date_sort_key
from ..pipeline import ClaimantInfo, MedicalEvent, PipelineContext

console = Console()


class ReportAgent(BaseAgent):
    name = "Report Agent"

    def _run(self, context: PipelineContext) -> PipelineContext:
        if context.claimant_info is None:
            raise RuntimeError("ReportAgent requires claimant_info in context.")

        claimant = context.claimant_info
        output_path = context.output_path
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if context.validation_issues:
            console.print(
                f"[yellow]Note:[/yellow] {len(context.validation_issues)} validation "
                f"issue(s) were corrected before report generation."
            )

        if context.layout_instruction:
            logger.info("Building custom layout report: %s", context.layout_instruction)
            console.print(
                f"Building custom layout report: [cyan]{context.layout_instruction}[/cyan]"
            )
            try:
                custom_rows = apply_custom_layout(
                    claimant,
                    context.layout_instruction,
                    model=context.model,
                )
                out_path = _generate_custom_report(
                    claimant, custom_rows, context.layout_instruction, output_path,
                    completion_through=context.completion_through,
                )
                logger.info("Custom report generated: %s (%d rows)", out_path, len(custom_rows))
            except Exception as exc:
                logger.warning(
                    "Custom layout LLM call failed (%s) — falling back to default template.", exc
                )
                console.print(
                    f"[yellow]Custom layout failed ({exc}) — "
                    f"falling back to default Word template layout.[/yellow]"
                )
                out_path = _generate_template_report(
                    claimant, context.template_path, output_path,
                    completion_through=context.completion_through,
                    medical_sections=context.medical_sections,
                )
                logger.info("Template report generated (fallback): %s", out_path)
        else:
            logger.info("Populating Word template: %s", context.template_path)
            console.print("Populating Word template…")
            out_path = _generate_template_report(
                claimant, context.template_path, output_path,
                completion_through=context.completion_through,
                medical_sections=context.medical_sections,
            )
            logger.info("Template report generated: %s", out_path)

        context.report_path = out_path
        console.print(f"[bold green]Report saved →[/bold green] [cyan]{out_path}[/cyan]")
        return context


# ---------------------------------------------------------------------------
# Exhibit ref resolution
# ---------------------------------------------------------------------------

def _parse_pdf_page(ref: str) -> int | None:
    """Extract the integer page number from a ref string like 'Pg 292' or 'Page 292'."""
    m = re.search(r"\d+", ref or "")
    return int(m.group()) if m else None


def _resolve_exhibit_ref(ref: str, medical_sections: list[dict]) -> str:
    """Convert a PDF-page ref (e.g. 'Pg 292') to an exhibit-size ref ('Pg 91').

    For each F-section, the convention is that ALL events from that exhibit share
    a single REF equal to the section's total page count (e.g. 'Pg 91' for a
    91-page exhibit). This matches standard SSA medical summary format.

    When consecutive sections overlap (a later cover page falls inside the previous
    section's page range), prefer the section whose cover page starts LATEST —
    that is the most specific / correct exhibit for the event.
    """
    if not medical_sections:
        return ref
    pdf_page = _parse_pdf_page(ref)
    if pdf_page is None:
        return ref

    matches = [
        s for s in medical_sections
        if s["start_page"] <= pdf_page <= s["end_page"]
    ]
    if not matches:
        return ref

    # When multiple sections share the page, the one starting latest is the correct exhibit
    best = max(matches, key=lambda s: s["start_page"])
    return f"Pg {best['total_pages']}"


# ---------------------------------------------------------------------------
# Template report
# ---------------------------------------------------------------------------

def _generate_template_report(
    claimant: ClaimantInfo,
    template_path: Path,
    output_path: Path,
    completion_through: str = "",
    medical_sections: list[dict] | None = None,
) -> Path:
    doc = Document(str(template_path))
    _populate_template(
        doc, claimant,
        completion_through=completion_through,
        medical_sections=medical_sections or [],
    )
    doc.save(str(output_path))
    return output_path


def _populate_template(
    doc: Document,
    claimant: ClaimantInfo,
    completion_through: str = "",
    medical_sections: list[dict] | None = None,
) -> None:
    sections = medical_sections or []
    last_updated = date.today().strftime("%A, %B %d, %Y")
    if completion_through:
        last_updated = f"{last_updated} \u2013 COMPLETED THROUGH {completion_through}"

    replacements = {
        "{{NAME}}": claimant.name,
        "{{SSN}}": claimant.ssn,
        "{{TITLE}}": claimant.title,
        "{{DLI}}": claimant.dli,
        "{{AOD}}": claimant.aod,
        "{{DOB}}": claimant.dob,
        "{{AGE_AT_AOD}}": claimant.age_at_aod,
        "{{CURRENT_AGE}}": claimant.current_age,
        "{{LAST_GRADE}}": claimant.last_grade,
        "{{SPECIAL_ED}}": claimant.special_ed,
        "{{LAST_UPDATED}}": last_updated,
        "{{IMPAIRMENTS}}": "; ".join(claimant.alleged_impairments),
    }

    # Replace placeholders in body paragraphs; also clear the decorative
    # "DATE  PROVIDER  REASON  REF" text paragraph that duplicates the table header.
    for para in doc.paragraphs:
        text_upper = para.text.upper()
        if "DATE" in text_upper and "PROVIDER" in text_upper and "REASON" in text_upper:
            # Clear all runs in this decorative header paragraph
            for run in para.runs:
                run.text = ""
            continue
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)
        headers = [cell.text.strip().upper() for cell in table.rows[0].cells]
        if "DATE" in headers:
            _fill_events_table(table, claimant.medical_events, sections)
            break


def _replace_in_paragraph(para, replacements: dict[str, str]) -> None:
    """Replace placeholder tokens in a paragraph, handling run-split tokens.

    Word often splits a token like {{NAME}} across multiple runs. We merge all
    runs into the first one, do the replacement, then zero out the rest so the
    paragraph structure is preserved without losing the containing paragraph.
    """
    full_text = "".join(run.text for run in para.runs)
    if not any(key in full_text for key in replacements):
        return

    for key, value in replacements.items():
        full_text = full_text.replace(key, value)

    for i, run in enumerate(para.runs):
        run.text = full_text if i == 0 else ""


def _fill_events_table(
    table,
    events: list[MedicalEvent],
    medical_sections: list[dict] = [],
) -> None:
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)

    # Sort chronologically with proper date parsing, then fill rows
    sorted_events = sorted(events, key=lambda e: (_date_sort_key(e.date), e.provider))

    for event in sorted_events:
        exhibit_ref = _resolve_exhibit_ref(event.ref, medical_sections)
        row = table.add_row()
        values = [event.date, event.provider, event.reason, exhibit_ref]
        for i, cell in enumerate(row.cells):
            cell.text = values[i] if i < len(values) else ""
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Custom layout report
# ---------------------------------------------------------------------------

def _generate_custom_report(
    claimant: ClaimantInfo,
    custom_rows: list[dict[str, str]],
    layout_instruction: str,
    output_path: Path,
    completion_through: str = "",
) -> Path:
    doc = Document()
    _add_header_block(doc, claimant, completion_through=completion_through)
    _add_custom_table(doc, custom_rows, layout_instruction)
    doc.save(str(output_path))
    return output_path


def _add_header_block(
    doc: Document,
    claimant: ClaimantInfo,
    completion_through: str = "",
) -> None:
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(0)
    run = title_para.add_run("MEDICAL SUMMARY")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"

    last_updated = date.today().strftime("%A, %B %d, %Y")
    if completion_through:
        last_updated = f"{last_updated} \u2013 COMPLETED THROUGH {completion_through}"
    for line in [
        f"RE: {claimant.name}    SSN: {claimant.ssn}    Title: {claimant.title}    DLI: {claimant.dli}",
        f"AOD: {claimant.aod}    Date of Birth: {claimant.dob}    "
        f"Age at AOD: {claimant.age_at_aod}    Current Age: {claimant.current_age}",
        f"Last Grade Completed: {claimant.last_grade}    Attended Special Ed Classes: {claimant.special_ed}",
        f"Last Updated: {last_updated}",
    ]:
        doc.add_paragraph(line)


def _add_custom_table(
    doc: Document,
    rows: list[dict[str, str]],
    layout_instruction: str,
) -> None:
    if not rows:
        doc.add_paragraph("No medical events found.")
        return

    columns = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col.replace("_", " ").title()

    for row_data in rows:
        row = table.add_row()
        for i, col in enumerate(columns):
            row.cells[i].text = row_data.get(col, "")
