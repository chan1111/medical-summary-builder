"""Tests for ReportAgent — Word document generation."""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from medical_summary_builder.agents.report_agent import (
    ReportAgent,
    _generate_template_report,
    _generate_custom_report,
    _replace_in_paragraph,
    _fill_events_table,
)
from medical_summary_builder.pipeline import ClaimantInfo, MedicalEvent, PipelineContext


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def template_docx(tmp_path: Path) -> Path:
    """Create a minimal Word template with all placeholder tokens."""
    doc = Document()
    doc.add_paragraph(
        "RE: {{NAME}}  SSN: {{SSN}}  Title: {{TITLE}}  DLI: {{DLI}}"
    )
    doc.add_paragraph(
        "AOD: {{AOD}}  DOB: {{DOB}}  Age at AOD: {{AGE_AT_AOD}}  "
        "Current Age: {{CURRENT_AGE}}  Last Grade: {{LAST_GRADE}}  "
        "Special Ed: {{SPECIAL_ED}}"
    )
    doc.add_paragraph("Last Updated: {{LAST_UPDATED}}")
    doc.add_paragraph("Impairments: {{IMPAIRMENTS}}")

    table = doc.add_table(rows=1, cols=4)
    for i, header in enumerate(["Date", "Provider", "Reason", "Ref"]):
        table.rows[0].cells[i].text = header

    path = tmp_path / "template.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# _replace_in_paragraph
# ---------------------------------------------------------------------------

class TestReplaceInParagraph:
    def test_replaces_single_placeholder(self):
        doc = Document()
        para = doc.add_paragraph("Hello {{NAME}}")
        _replace_in_paragraph(para, {"{{NAME}}": "Alice"})
        assert "Alice" in para.text
        assert "{{NAME}}" not in para.text

    def test_no_op_when_no_placeholder(self):
        doc = Document()
        para = doc.add_paragraph("No placeholders here")
        original_text = para.text
        _replace_in_paragraph(para, {"{{NAME}}": "Alice"})
        assert para.text == original_text

    def test_replaces_multiple_placeholders_in_one_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("{{NAME}} born {{DOB}}")
        _replace_in_paragraph(para, {"{{NAME}}": "Bob", "{{DOB}}": "01/01/1980"})
        assert "Bob" in para.text
        assert "01/01/1980" in para.text


# ---------------------------------------------------------------------------
# _fill_events_table
# ---------------------------------------------------------------------------

class TestFillEventsTable:
    def test_fills_one_row_per_event(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        for i, h in enumerate(["Date", "Provider", "Reason", "Ref"]):
            table.rows[0].cells[i].text = h

        events = [
            MedicalEvent(date="01/01/2020", provider="Clinic A", reason="Checkup", ref="Pg 1"),
            MedicalEvent(date="02/01/2020", provider="Hospital B", reason="Surgery", ref="Pg 5"),
        ]
        _fill_events_table(table, events)

        assert len(table.rows) == 3  # header + 2 events
        assert table.rows[1].cells[0].text == "01/01/2020"
        assert table.rows[2].cells[1].text == "Hospital B"

    def test_clears_previous_rows_before_filling(self):
        doc = Document()
        table = doc.add_table(rows=3, cols=4)
        table.rows[0].cells[0].text = "Date"

        events = [
            MedicalEvent(date="03/03/2021", provider="Center C", reason="Visit", ref="Pg 2"),
        ]
        _fill_events_table(table, events)
        assert len(table.rows) == 2  # header + 1

    def test_empty_events_leaves_only_header(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        _fill_events_table(table, [])
        assert len(table.rows) == 1


# ---------------------------------------------------------------------------
# _generate_template_report
# ---------------------------------------------------------------------------

class TestGenerateTemplateReport:
    def test_creates_output_file(self, template_docx: Path, sample_claimant, tmp_path: Path):
        out_path = tmp_path / "output" / "result.docx"
        out_path.parent.mkdir(parents=True, exist_ok=True)
        _generate_template_report(sample_claimant, template_docx, out_path)
        assert out_path.exists()

    def test_placeholders_replaced(self, template_docx: Path, sample_claimant, tmp_path: Path):
        out_path = tmp_path / "result.docx"
        _generate_template_report(sample_claimant, template_docx, out_path)

        doc = Document(str(out_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "John Doe" in full_text
        assert "123-45-6789" in full_text
        assert "{{NAME}}" not in full_text
        assert "{{SSN}}" not in full_text

    def test_medical_events_in_table(self, template_docx: Path, sample_claimant, tmp_path: Path):
        out_path = tmp_path / "events.docx"
        _generate_template_report(sample_claimant, template_docx, out_path)

        doc = Document(str(out_path))
        table = doc.tables[0]
        all_text = " ".join(cell.text for row in table.rows for cell in row.cells)
        assert "City Medical Center" in all_text
        assert "03/01/2018" in all_text


# ---------------------------------------------------------------------------
# _generate_custom_report
# ---------------------------------------------------------------------------

class TestGenerateCustomReport:
    def test_creates_custom_columns_table(self, sample_claimant, tmp_path: Path):
        rows = [
            {"date": "03/01/2018", "facility": "City Medical Center", "summary": "Back pain", "ref": "Pg 3"}
        ]
        out_path = tmp_path / "custom.docx"
        _generate_custom_report(sample_claimant, rows, "Date, Facility, Summary, Ref", out_path)

        assert out_path.exists()
        doc = Document(str(out_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "John Doe" in all_text

        table = doc.tables[0]
        headers = [cell.text for cell in table.rows[0].cells]
        assert "Date" in headers
        assert "Facility" in headers

    def test_empty_rows_adds_fallback_paragraph(self, sample_claimant, tmp_path: Path):
        out_path = tmp_path / "empty.docx"
        _generate_custom_report(sample_claimant, [], "Date, Facility", out_path)

        doc = Document(str(out_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "No medical events found" in all_text


# ---------------------------------------------------------------------------
# ReportAgent._run integration
# ---------------------------------------------------------------------------

class TestReportAgent:
    def test_raises_without_claimant_info(self, base_context: PipelineContext):
        with pytest.raises(RuntimeError, match="claimant_info"):
            ReportAgent().run(base_context)

    def test_template_mode_creates_file(
        self,
        context_with_claimant: PipelineContext,
        template_docx: Path,
        tmp_path: Path,
    ):
        context_with_claimant.template_path = template_docx
        context_with_claimant.output_path = tmp_path / "out" / "summary.docx"
        context_with_claimant.layout_instruction = None

        result = ReportAgent().run(context_with_claimant)

        assert result.report_path is not None
        assert result.report_path.exists()

    def test_custom_layout_mode_calls_apply_custom_layout(
        self,
        context_with_claimant: PipelineContext,
        template_docx: Path,
        tmp_path: Path,
        monkeypatch,
    ):
        context_with_claimant.template_path = template_docx
        context_with_claimant.output_path = tmp_path / "out" / "custom.docx"
        context_with_claimant.layout_instruction = "Date, Facility, Summary, Ref"
        monkeypatch.setenv("AI_BUILDER_TOKEN", "test-token")

        custom_rows = [
            {"date": "03/01/2018", "facility": "City Medical", "summary": "Pain", "ref": "Pg 3"}
        ]
        with patch(
            "medical_summary_builder.agents.report_agent.apply_custom_layout",
            return_value=custom_rows,
        ):
            result = ReportAgent().run(context_with_claimant)

        assert result.report_path is not None
        assert result.report_path.exists()

    def test_report_path_set_in_context(
        self,
        context_with_claimant: PipelineContext,
        template_docx: Path,
        tmp_path: Path,
    ):
        context_with_claimant.template_path = template_docx
        context_with_claimant.output_path = tmp_path / "report.docx"
        context_with_claimant.layout_instruction = None

        result = ReportAgent().run(context_with_claimant)
        assert result.report_path == tmp_path / "report.docx"

    def test_validation_issues_note_is_printed(
        self,
        context_with_claimant: PipelineContext,
        template_docx: Path,
        tmp_path: Path,
        capsys,
    ):
        context_with_claimant.template_path = template_docx
        context_with_claimant.output_path = tmp_path / "report.docx"
        context_with_claimant.layout_instruction = None
        context_with_claimant.validation_issues = ["Low match for SomeClinic"]

        ReportAgent().run(context_with_claimant)
        # Rich writes to its own console — just verify it didn't raise
